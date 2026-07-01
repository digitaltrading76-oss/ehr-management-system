from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import datetime
import re

LABOR_REFERENCES = [
    {
        "trigger": ["theft", "fraud", "falsified", "dishonesty", "cash", "remittance", "stolen", "missing parcel"],
        "title": "Labor Code, Article 297 - Just causes; serious misconduct, fraud, or willful breach of trust may be relevant depending on evidence.",
        "quote": "Article 297 recognizes just causes including serious misconduct, willful disobedience, gross and habitual neglect of duties, fraud or willful breach of trust, commission of a crime against the employer or immediate family, and analogous causes.",
        "caution": "Application depends on substantial evidence and completion of the required employee notice and opportunity to respond."
    },
    {
        "trigger": ["refuse", "disobey", "insubordination"],
        "title": "Labor Code, Article 297 - Willful disobedience may be considered if supported by facts.",
        "quote": "A lawful and reasonable work-related order must be identified, and the refusal must be shown by evidence.",
        "caution": "The order, knowledge of the order, refusal, and employee side should be documented."
    },
    {
        "trigger": ["negligence", "on hold", "parcel", "gps", "timestamp", "damage", "delivery"],
        "title": "Labor Code, Article 297 - Gross and habitual neglect of duties may be relevant only when supported by evidence.",
        "quote": "Neglect-based action should be evaluated against the documents, prior record, seriousness, and proportionality of the intended action.",
        "caution": "A single operational lapse may require further evaluation before any serious disciplinary action."
    },
    {
        "trigger": ["salary", "wage", "payroll", "deduction", "underpaid", "unpaid"],
        "title": "Labor standards review - wage or payroll concerns should first be validated as compliance matters.",
        "quote": "Payroll-related concerns require verification of computation, attendance, deductions, and proof of payment.",
        "caution": "This may not be a worker violation. Treat first as payroll/compliance validation."
    },
    {
        "trigger": ["accident", "injury", "medical", "hospital", "crash"],
        "title": "Safety and welfare review - accident cases require documentation and assistance verification.",
        "quote": "Accident and medical records should be gathered before determining whether any misconduct or negligence exists.",
        "caution": "Prioritize worker welfare and safety documentation."
    }
]

DOCUMENT_KEYWORDS = {
    "Incident Report": ["incident", "ir", "report"],
    "Employee Explanation": ["explanation", "salaysay", "explain", "statement"],
    "Witness Statement": ["witness", "testimony", "affidavit"],
    "GPS / System Logs": ["gps", "log", "timestamp", "tracking"],
    "Payroll / Attendance": ["payroll", "salary", "attendance", "timesheet", "dtr"],
    "Photos / Screenshots": ["photo", "image", "screenshot", "jpg", "jpeg", "png"],
    "CCTV / Video": ["cctv", "video", "mp4", "mov"],
    "Medical / Accident Docs": ["medical", "hospital", "accident", "injury"]
}

def _contains(text, words):
    t = text.lower()
    return any(w in t for w in words)

def classify_documents(files):
    results = []
    for label, keys in DOCUMENT_KEYWORDS.items():
        matches = [f for f in files if _contains(f, keys)]
        results.append({
            "document": label,
            "status": "Submitted" if matches else "Missing / Not Uploaded",
            "remarks": ", ".join(matches) if matches else "Not found in submitted files"
        })
    return results

def labor_references_for_case(case):
    text = (case.get("incident_category","") + " " + case.get("extracted_text","") + " " + " ".join(case.get("uploaded_files",[]))).lower()
    refs = [r for r in LABOR_REFERENCES if _contains(text, r["trigger"])]
    if not refs:
        refs = [{
            "title": "General Philippine labor standards review",
            "quote": "No specific labor provision was automatically matched from the uploaded documents.",
            "caution": "Apply company rules first, then verify compliance with employee notice, opportunity to respond, and documentation requirements."
        }]
    return refs

def extract_facts(case):
    extracted = case.get("extracted_text","").strip()
    files = case.get("uploaded_files",[])
    facts = [
        f"Case was submitted by {case.get('submitted_by','')} from {case.get('area','')}.",
        f"Worker identified in the case record: {case.get('worker_name','')} ({case.get('worker_type','')}).",
        f"Incident category selected during submission: {case.get('incident_category','')}."
    ]
    facts.append(f"{len(files)} file(s) were uploaded for review." if files else "No supporting file was uploaded.")
    if extracted:
        preview = re.sub(r"\s+", " ", extracted)[:600]
        facts.append("Extracted content preview from uploaded files: " + preview)
    return facts

def evidence_level(case):
    files = case.get("uploaded_files",[])
    assessment = case.get("assessment",{})
    scores = assessment.get("scores",{})
    evidence = int(scores.get("evidence_sufficiency",0))
    if not files:
        return "Insufficient Evidence", "No file has been uploaded."
    if evidence >= 75:
        return "Strong / Substantial File Support", "Several supporting documents are available. HR should still verify authenticity, relevance, and employee response."
    if evidence >= 50:
        return "Moderate Evidence", "There are supporting documents, but additional records or explanation may still be needed."
    return "Limited Evidence", "The current file set is incomplete. Additional documents are recommended before further action."

def recommended_actions(case):
    assessment = case.get("assessment",{})
    next_move = assessment.get("recommended_next_move","")
    missing = assessment.get("missing_requirements",[])
    actions = []
    if missing:
        actions.append("Request additional documents from the coordinator.")
    if "payroll" in next_move.lower() or "salary" in next_move.lower():
        actions.append("Request payroll, attendance, payout proof, and deduction basis.")
    if "high-risk" in next_move.lower():
        actions.append("Secure evidence chain and employee explanation before notice/decision stage.")
    if "nte" in next_move.lower() or "notice" in next_move.lower():
        actions.append("Prepare Notice to Explain only after HR reviewer confirms file sufficiency.")
    if not actions:
        actions.append("Continue HR review and determine if additional documents, employee explanation, or closure is appropriate.")
    return actions

def build_assessment_context(case):
    assessment = case.get("assessment",{})
    scores = assessment.get("scores",{})
    return {
        "case_id": case.get("case_id",""),
        "date_generated": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
        "case_info": {
            "Worker Name": case.get("worker_name",""),
            "Worker Type": case.get("worker_type",""),
            "Coordinator": case.get("submitted_by",""),
            "Area": case.get("area",""),
            "Incident Category": case.get("incident_category",""),
            "Current Status": case.get("status",""),
        },
        "documents": classify_documents(case.get("uploaded_files",[])),
        "files": case.get("uploaded_files",[]),
        "facts": extract_facts(case),
        "policy": assessment.get("policy_assessment",[]),
        "labor_refs": labor_references_for_case(case),
        "checklist": assessment.get("due_process_checklist",[]),
        "missing": assessment.get("missing_requirements",[]),
        "scores": scores,
        "evidence_level": evidence_level(case),
        "recommended_next_move": assessment.get("recommended_next_move",""),
        "recommended_actions": recommended_actions(case),
        "safeguard": assessment.get("safeguard","No final HR action should be recommended until the worker has been properly notified and given an opportunity to respond."),
        "extracted_preview": assessment.get("extracted_content_preview", case.get("extracted_text",""))[:2500],
    }

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "D9EAF7")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    document.add_paragraph("")
    return table

def create_docx_report(case, output_path):
    ctx = build_assessment_context(case)
    doc = Document()

    sec = doc.sections[0]
    sec.top_margin = Inches(0.6)
    sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("EHR CASE ASSESSMENT REPORT")
    r.bold = True
    r.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Embedded Human Resource • Powered by Conglomerate Corp").italic = True

    doc.add_paragraph(f"Case Number: {ctx['case_id']}")
    doc.add_paragraph(f"Date Generated: {ctx['date_generated']}")

    doc.add_heading("1. Case Information", level=1)
    add_table(doc, ["Field", "Details"], [[k,v] for k,v in ctx["case_info"].items()])

    doc.add_heading("2. Documents Received", level=1)
    add_table(doc, ["Document Type", "Status", "Remarks"], [[d["document"], d["status"], d["remarks"]] for d in ctx["documents"]])

    doc.add_heading("3. Uploaded Files", level=1)
    if ctx["files"]:
        for f in ctx["files"]:
            doc.add_paragraph(f)
    else:
        doc.add_paragraph("No uploaded file listed.")

    doc.add_heading("4. Facts Established From Uploaded Documents", level=1)
    for fact in ctx["facts"]:
        doc.add_paragraph(fact)

    doc.add_heading("5. Company Code of Conduct Review", level=1)
    policy_rows = [[p.get("policy_category",""), p.get("possible_violation",""), p.get("policy_basis",""), p.get("usual_next_step","")] for p in ctx["policy"]]
    add_table(doc, ["Policy Category", "Possible Violation", "Basis", "Usual Next Step"], policy_rows)

    doc.add_heading("6. Philippine Labor Standards Reference", level=1)
    labor_rows = [[r.get("title",""), r.get("quote",""), r.get("caution","")] for r in ctx["labor_refs"]]
    add_table(doc, ["Reference", "Relevant Text / Principle", "Caution"], labor_rows)

    doc.add_heading("7. Review Checklist", level=1)
    add_table(doc, ["Requirement", "Status"], [[x.get("item",""), "Complete" if x.get("status") else "Missing / Pending"] for x in ctx["checklist"]])

    doc.add_heading("8. Missing Documents / Follow-up Needed", level=1)
    if ctx["missing"]:
        for m in ctx["missing"]:
            doc.add_paragraph(m)
    else:
        doc.add_paragraph("No missing requirement identified.")

    doc.add_heading("9. Evidence Assessment", level=1)
    level, explanation = ctx["evidence_level"]
    doc.add_paragraph(f"Evidence Level: {level}")
    doc.add_paragraph(explanation)

    doc.add_heading("10. Preliminary HR Assessment", level=1)
    doc.add_paragraph(ctx["recommended_next_move"])

    doc.add_heading("11. Recommended Next Actions", level=1)
    for a in ctx["recommended_actions"]:
        doc.add_paragraph(a)

    doc.add_heading("12. Coordinator Action Required", level=1)
    for m in ctx["missing"] or ["Await further instruction from EHR reviewer."]:
        doc.add_paragraph(m)

    doc.add_heading("13. HR Reviewer Notes", level=1)
    doc.add_paragraph("Reviewer Notes:")
    doc.add_paragraph("\n\n\n")

    doc.add_heading("14. Safeguard", level=1)
    doc.add_paragraph(ctx["safeguard"])

    doc.add_heading("15. Extracted Content Preview", level=1)
    doc.add_paragraph(ctx["extracted_preview"] or "No extracted content available.")

    doc.save(output_path)
    return output_path
